import os
import google_auth_oauthlib.flow
import googleapiclient.discovery
import googleapiclient.errors
import urllib.request
import re
from PIL import Image
from docx import Document
from docx.shared import Inches
from datetime import datetime
import hashlib

# Set up API credentials
api_service_name = "youtube"
api_version = "v3"
client_secrets_file = "S:\projects\pp extractor yt\client_secret_181788042811-adt0p4q6bg9nacsu25cb29fn8erjsfin.apps.googleusercontent.com.json"  # Path to your client secret file
scopes = ["https://www.googleapis.com/auth/youtube.readonly"]

# Create the YouTube API client
def create_youtube_client():
    flow = google_auth_oauthlib.flow.InstalledAppFlow.from_client_secrets_file(
        client_secrets_file, scopes)
    credentials = flow.run_local_server(port=8080)  # Specify a port
    return googleapiclient.discovery.build(
        api_service_name, api_version, credentials=credentials)

# Get subscribed channels
def get_subscribed_channels(youtube):
    channels = []
    request = youtube.subscriptions().list(
        part="snippet",
        mine=True
    )
    while request is not None:
        response = request.execute()
        channels.extend([item['snippet']['resourceId']['channelId'] for item in response.get('items', [])])
        request = youtube.subscriptions().list_next(request, response)
    return channels

# Download profile pictures and convert to JPG
def download_and_convert_profile_pictures(youtube, channels, output_folder, word_file):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    doc = Document()
    doc.add_heading('Subscribed YouTube Channels', level=1)

    for channel_id in channels:
        try:
            channel_info = youtube.channels().list(
                part="snippet,statistics",
                id=channel_id
            ).execute().get('items', [])[0]

            snippet = channel_info.get('snippet', {})
            statistics = channel_info.get('statistics', {})
            channel_name = snippet.get('title')

            # Replace problematic characters in channel name
            channel_name = re.sub(r'[\\/:"*?<>|]', '', channel_name)

            # Skip channels with names exceeding 100 characters
            if len(channel_name) > 100:
                print(f"Skipping channel '{channel_name}' as its name exceeds 100 characters.")
                continue

            channel_link = f"https://www.youtube.com/channel/{channel_id}"
            subscription_date = datetime.now().strftime('%Y-%m-%d')

            doc.add_heading(channel_name, level=2)
            doc.add_paragraph(channel_link)

            profile_picture_url = snippet.get('thumbnails', {}).get('default', {}).get('url')
            if profile_picture_url:
                # Generate filename using channel name
                file_name = f"{channel_name}.jpg"

                output_path = os.path.join(output_folder, file_name)

                urllib.request.urlretrieve(profile_picture_url, output_path)

                img = Image.open(output_path)
                img.convert("RGB").save(output_path)

                doc.add_picture(output_path, width=Inches(1.0))
        except Exception as e:
            print(f"Error processing channel {channel_name}: {e}")

        doc.add_paragraph('')

    doc.save(word_file)


# Main function
def main():
    youtube = create_youtube_client()
    channels = get_subscribed_channels(youtube)
    output_folder = "profile_pictures"
    word_file = "subscribed_channels.docx"
    download_and_convert_profile_pictures(youtube, channels, output_folder, word_file)
    
    # Delete images without extension
    delete_images_without_extension(output_folder)

# Function to delete images without extension
def delete_images_without_extension(folder):
    for filename in os.listdir(folder):
        if os.path.isfile(os.path.join(folder, filename)) and '.' not in filename:
            os.remove(os.path.join(folder, filename))

if __name__ == "__main__":
    main()
